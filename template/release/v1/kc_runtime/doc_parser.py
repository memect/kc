"""
Minimal document parser for the release runtime.

Strategy: try native Python parsers first (pypdf, python-docx),
fall back to LibreOffice CLI if natives unavailable AND lo is on
PATH, finally fall back to UTF-8 plaintext read. Each strategy
records what it tried via the result dict so workflows can decide
whether to trust the text.

This is a release-time helper — KC's CLI mode uses its own document
parsing pipeline (src/agent/document-parser.js + LibreOffice).
"""

from __future__ import annotations

import json
import shutil
import subprocess
from pathlib import Path


def preflight(doc: Path) -> dict:
    """
    Verify a document is parseable; return a small status dict.
    Lets workflows skip cleanly when the parse will fail rather than
    burning a worker-LLM call on an unreadable file.
    """
    if not doc.exists():
        return {"ok": False, "reason": "not_found", "path": str(doc)}
    if not doc.is_file():
        return {"ok": False, "reason": "not_file", "path": str(doc)}
    if doc.stat().st_size == 0:
        return {"ok": False, "reason": "empty", "path": str(doc)}
    return {"ok": True, "path": str(doc), "size_bytes": doc.stat().st_size}


def extract_text(doc: Path) -> dict:
    """
    Pull text out of a document. Returns:
      { "text": "...", "via": "<strategy>", "ok": bool, "error"?: str }
    Strategies tried in order:
      1. Suffix-specific native parser (pypdf for .pdf, python-docx for .docx)
      2. LibreOffice CLI (`soffice --headless --convert-to txt`) if on PATH
      3. UTF-8 plaintext (.txt, .md, or any file with text-like bytes)
    """
    suffix = doc.suffix.lower()

    if suffix == ".pdf":
        text = _try_pypdf(doc)
        if text is not None:
            return {"text": text, "via": "pypdf", "ok": True}

    if suffix in (".docx",):
        text = _try_python_docx(doc)
        if text is not None:
            return {"text": text, "via": "python-docx", "ok": True}

    # LibreOffice fallback for anything we couldn't parse natively
    if suffix in (".pdf", ".doc", ".docx", ".odt", ".rtf"):
        text = _try_libreoffice(doc)
        if text is not None:
            return {"text": text, "via": "libreoffice", "ok": True}

    # Plaintext fallback (covers .txt, .md, .csv, .json, etc.)
    try:
        text = doc.read_text(encoding="utf-8")
        return {"text": text, "via": "plaintext_utf8", "ok": True}
    except UnicodeDecodeError:
        try:
            text = doc.read_text(encoding="gbk")  # common in Chinese corpora
            return {"text": text, "via": "plaintext_gbk", "ok": True}
        except Exception as exc:
            return {"text": "", "via": "none", "ok": False, "error": str(exc)}


# --- internals ---


def _try_pypdf(doc: Path):
    try:
        import pypdf  # type: ignore
    except ImportError:
        return None
    try:
        reader = pypdf.PdfReader(str(doc))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return None


def _try_python_docx(doc: Path):
    try:
        import docx  # python-docx
    except ImportError:
        return None
    try:
        d = docx.Document(str(doc))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return None


def _try_libreoffice(doc: Path):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    out_dir = doc.parent / ".kc-lo-out"
    out_dir.mkdir(exist_ok=True)
    try:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "txt", "--outdir", str(out_dir), str(doc)],
            capture_output=True,
            timeout=60,
            check=True,
        )
        txt_path = out_dir / (doc.stem + ".txt")
        if txt_path.exists():
            return txt_path.read_text(encoding="utf-8")
    except Exception:
        return None
    return None
